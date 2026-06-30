from __future__ import annotations

import re
import tempfile
from collections import Counter
from pathlib import Path
from typing import BinaryIO, Dict, List, Tuple

import pandas as pd


SKIP_COLUMN_PATTERNS = [
    "sample",
    "flag",
    "count",
    "terminate",
    "elapsed",
    "starttime",
    "endtime",
    "password",
    "unique",
    "respid",
    "respnum",
    "serial",
    "token",
    "email",
    "phone",
    "mobile",
    "name",
    "address",
    "ip",
    "url",
    "browser",
    "device",
    "dummy",
    "hidden question",
    "punching",
    "none of the above",
    "are you employed in the following",
    "screenout",
    "screener",
    "would you like to be contacted",
    "keep your feedback anonymous",
    "anonymous",
    "prior permission",
]

CATEGORICAL_LABEL_PATTERNS = [
    "employment status",
    "current primary employment",
    "gender",
    "age",
    "region",
    "vertical",
    "profile",
    "role",
    "department",
    "designation",
    "category",
    "type",
    "tenure",
    "which of the following best describes",
]

RATING_LABEL_PATTERNS = [
    "rate",
    "rating",
    "satisfaction",
    "quality",
    "value",
    "ease",
    "performance",
    "scalability",
    "security",
    "reliability",
    "flexibility",
    "improvement",
    "implementation",
    "support",
    "relationship",
    "communication",
    "recognition",
    "preference",
]

NPS_LABEL_PATTERNS = [
    "recommend",
    "nps",
    "advocacy",
    "confident recommending",
    "colleague",
    "peer",
]


def read_survey_file(uploaded_file: BinaryIO) -> Tuple[pd.DataFrame, Dict[str, str]]:
    name = uploaded_file.name.lower()
    labels: Dict[str, str] = {}

    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file), labels

    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded_file), labels

    if name.endswith(".sav"):
        try:
            import pyreadstat
        except ImportError as exc:
            raise RuntimeError("SPSS .sav support needs pyreadstat in requirements.txt.") from exc

        with tempfile.NamedTemporaryFile(delete=False, suffix=".sav") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        try:
            df, meta = pyreadstat.read_sav(tmp_path, apply_value_formats=False)
            labels = meta.column_names_to_labels or {}
            return df, labels
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    raise ValueError("Please upload a CSV, Excel, or SPSS .sav file.")


def extract_questionnaire_text(uploaded_file: BinaryIO | None) -> str:
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

    if name.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX questionnaire support needs python-docx in requirements.txt.") from exc

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        try:
            doc = Document(tmp_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    return uploaded_file.getvalue().decode("utf-8", errors="ignore")


def readable_label(column: str, labels: Dict[str, str], questionnaire_text: str) -> str:
    if labels.get(column):
        return labels[column]

    cleaned = column.replace("_", " ").strip()
    if questionnaire_text and cleaned.lower() in questionnaire_text.lower():
        return cleaned
    return cleaned


def should_skip_column(column: str, label: str) -> bool:
    text = f"{column} {label}".lower()
    if any(pattern in text for pattern in SKIP_COLUMN_PATTERNS):
        return True
    if any(pattern in text for pattern in CATEGORICAL_LABEL_PATTERNS):
        return True
    return False


def contains_pattern(text: str, patterns: List[str]) -> bool:
    lowered = text.lower()
    return any(pattern in lowered for pattern in patterns)


def concise_theme(label: str, prefix: str) -> str:
    text = extract_item_name(label)
    lowered = text.lower()

    replacements = [
        ("overall quality of products and services", "Overall Quality"),
        ("overall quality", "Overall Quality"),
        ("value in terms of business impact", "Business Impact"),
        ("business impact", "Business Impact"),
        ("ease of doing business", "Ease of Doing Business"),
        ("confident recommending", "Recommendation Advocacy"),
        ("recommend", "Recommendation Advocacy"),
        ("technology partner", "Technology Partner Rating"),
        ("preference", "Partner Preference"),
        ("intuitiveness", "User Experience"),
        ("user interface", "User Experience"),
        ("performance", "Performance"),
        ("scalability", "Scalability"),
        ("security", "Security"),
        ("reliability", "Reliability"),
        ("flexibility", "Flexibility"),
        ("continuous improvement", "Continuous Improvement"),
        ("rapid application development", "Low-Code Development"),
        ("low code", "Low-Code Development"),
        ("implementation and support", "Implementation Support"),
        ("implementation", "Implementation Support"),
        ("support", "Support Experience"),
        ("service level", "Support Responsiveness"),
        ("response time", "Support Responsiveness"),
        ("customer relationships", "Relationship Engagement"),
        ("long-term customer relationships", "Relationship Engagement"),
        ("customer meets", "Customer Engagement"),
        ("feedback forums", "Customer Engagement"),
        ("communication", "Communication"),
        ("product updates", "Communication"),
        ("recognition", "Project Team Recognition"),
        ("reduce costs", "Efficiency and Cost Reduction"),
        ("employee efficiency", "Efficiency and Cost Reduction"),
        ("regulatory", "Risk and Compliance"),
        ("comply", "Risk and Compliance"),
        ("customer experience", "Customer Experience"),
        ("remote working", "Remote Working"),
        ("return on investments", "ROI"),
        ("serve you better", "Open-Ended Improvement"),
        ("focus areas", "Future Focus Areas"),
        ("digital transformation", "Digital Transformation Journey"),
    ]

    for pattern, theme in replacements:
        if pattern in lowered:
            return theme

    cleaned = re.sub(r"please\s+rate\s+", "", text, flags=re.I)
    cleaned = re.sub(r"newgen('?s)?\s+", "", cleaned, flags=re.I)
    cleaned = re.sub(r"on the following parameters:?", "", cleaned, flags=re.I)
    cleaned = re.sub(r"which of the following\s+", "", cleaned, flags=re.I)
    cleaned = re.sub(r"::.*", "", cleaned).strip(" :-")

    words = re.findall(r"[A-Za-z0-9]+", cleaned)
    if not words:
        return prefix

    short = " ".join(words[:5]).title()
    return short if short else prefix


def extract_item_name(label: str) -> str:
    text = re.sub(r"\s+", " ", str(label)).strip()

    if ":" in text:
        after_colon = text.split(":", 1)[1].strip()
        if after_colon:
            text = after_colon

    for separator in [" - ", " – ", " — ", " :: "]:
        if separator in text:
            text = text.split(separator, 1)[0].strip()
            break

    text = re.sub(r"\([^)]*\)", "", text).strip()
    text = re.sub(r"^please\s+rate\s+", "", text, flags=re.I)
    text = re.sub(r"^newgen('?s)?\s+", "", text, flags=re.I)
    text = re.sub(r"on the following parameters:?", "", text, flags=re.I)
    text = re.sub(r"which of the following\s+", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip(" :-")

    words = text.split()
    if len(words) > 8:
        text = " ".join(words[:8])

    return text or str(label).strip()


def numeric_series(series: pd.Series, drop_cant_say: bool = False) -> pd.Series:
    values = pd.to_numeric(series, errors="coerce").dropna()
    if drop_cant_say:
        return values[values != 9]
    return values


def is_rating_1_to_5(values: pd.Series, label: str) -> bool:
    if values.empty:
        return False
    if not contains_pattern(label, RATING_LABEL_PATTERNS):
        return False
    if values.max() <= 2:
        return False
    return values.min() >= 1 and values.max() <= 5 and values.nunique() >= 2


def is_nps_0_to_10(values: pd.Series, label: str) -> bool:
    if values.empty:
        return False
    if not contains_pattern(label, NPS_LABEL_PATTERNS):
        return False
    return values.min() >= 0 and values.max() <= 10 and values.nunique() >= 5


def generate_rating_insight(insight_id: str, column: str, label: str, values: pd.Series) -> Dict[str, str]:
    mean = values.mean()
    top2 = (values >= 4).mean() * 100
    low = (values <= 2).mean() * 100
    theme = concise_theme(label, "Rating")

    if top2 >= 65 and low < 15:
        text = f"{theme} appears to be a relative strength, with {top2:.1f}% top-two-box ratings and a mean score of {mean:.2f}/5."
    elif low >= 20 or top2 < 45:
        text = f"{theme} needs human attention, as {low:.1f}% of valid respondents gave low ratings despite a mean score of {mean:.2f}/5."
    else:
        text = f"{theme} shows moderate customer confidence, with a mean score of {mean:.2f}/5 and {top2:.1f}% top-two-box ratings."

    return {
        "insight_id": insight_id,
        "theme": theme,
        "insight_text": text,
        "evidence_note": f"{column}: n={len(values)}, mean={mean:.2f}/5, top-two-box={top2:.1f}%, low ratings={low:.1f}%.",
    }


def generate_nps_insight(insight_id: str, column: str, label: str, values: pd.Series) -> Dict[str, str]:
    promoters = (values >= 9).mean() * 100
    passives = ((values >= 7) & (values <= 8)).mean() * 100
    detractors = (values <= 6).mean() * 100
    nps = promoters - detractors
    theme = concise_theme(label, "Recommendation Advocacy")

    if nps >= 40 and detractors < 15:
        text = f"{theme} shows strong customer advocacy, with an NPS-style score of {nps:.1f} and {promoters:.1f}% promoters."
    elif nps < 10 or detractors >= 30:
        text = f"{theme} needs human attention because advocacy is weak, with an NPS-style score of {nps:.1f} and {detractors:.1f}% detractors."
    else:
        text = f"{theme} shows moderate advocacy, with an NPS-style score of {nps:.1f}, {promoters:.1f}% promoters, {passives:.1f}% passives, and {detractors:.1f}% detractors."

    return {
        "insight_id": insight_id,
        "theme": theme,
        "insight_text": text,
        "evidence_note": f"{column}: n={len(values)}, promoters={promoters:.1f}%, passives={passives:.1f}%, detractors={detractors:.1f}%, NPS-style score={nps:.1f}.",
    }


def generate_binary_insight(insight_id: str, column: str, label: str, values: pd.Series, total_n: int) -> Dict[str, str] | None:
    unique = set(values.dropna().unique().tolist())
    if not unique.issubset({0, 1}) and not unique.issubset({1, 2}):
        return None

    selected = (values == 1).sum()
    pct = selected / total_n * 100 if total_n else 0
    if pct < 10:
        return None
    theme = concise_theme(label, "Selection Theme")
    lowered = label.lower()
    if "currently done" in lowered:
        insight_text = f"{theme} is currently done by {pct:.1f}% of respondents, making it a visible activity area in the study."
    elif pct >= 50:
        insight_text = f"{theme} is selected by a majority of respondents at {pct:.1f}%, making it a prominent survey theme."
    else:
        insight_text = f"{theme} is selected by {pct:.1f}% of respondents, making it a visible theme in the survey response pattern."

    return {
        "insight_id": insight_id,
        "theme": theme,
        "insight_text": insight_text,
        "evidence_note": f"{column}: selected n={int(selected)} out of total n={total_n}, selected percentage={pct:.1f}%.",
    }


def generate_text_insight(insight_id: str, column: str, label: str, series: pd.Series) -> Dict[str, str] | None:
    responses = [str(value).strip() for value in series.dropna() if str(value).strip()]
    if len(responses) < 10:
        return None

    stop_words = set(
        "the and for with this that have has had are was were you your our from product service services support customer customers solution solutions team very good great more need needs should can better improve improvement in on to of a an is it as by be we they their at all also".split()
    )
    words: List[str] = []
    for response in responses:
        for word in re.findall(r"[A-Za-z][A-Za-z-]{2,}", response.lower()):
            if word not in stop_words:
                words.append(word)

    top_words = [word for word, _ in Counter(words).most_common(6)]
    if not top_words:
        return None
    theme = concise_theme(label, "Open-Ended Theme")

    return {
        "insight_id": insight_id,
        "theme": theme,
        "insight_text": f"Open-ended responses for {theme} suggest recurring themes around {', '.join(top_words[:5])}; this should be coded qualitatively before client reporting.",
        "evidence_note": f"{column}: {len(responses)} open-ended responses reviewed; frequent terms include {', '.join(top_words)}.",
    }


def generate_insights(
    df: pd.DataFrame,
    labels: Dict[str, str] | None = None,
    questionnaire_text: str = "",
    max_insights: int = 25,
) -> pd.DataFrame:
    labels = labels or {}
    insights: List[Dict[str, str]] = []
    total_n = len(df)

    for column in df.columns:
        if len(insights) >= max_insights:
            break

        label = readable_label(str(column), labels, questionnaire_text)
        if should_skip_column(str(column), label):
            continue

        values = numeric_series(df[column])

        if len(values) >= max(10, total_n * 0.20):
            binary = generate_binary_insight(f"AI-{len(insights) + 1:03d}", str(column), label, values, total_n)
            if binary:
                insights.append(binary)
                continue
            if is_nps_0_to_10(values, label):
                insights.append(generate_nps_insight(f"AI-{len(insights) + 1:03d}", str(column), label, values))
                continue
            rating_values = numeric_series(df[column], drop_cant_say=True)
            if is_rating_1_to_5(rating_values, label):
                insights.append(generate_rating_insight(f"AI-{len(insights) + 1:03d}", str(column), label, rating_values))
                continue

        if df[column].dtype == "object":
            text_insight = generate_text_insight(f"AI-{len(insights) + 1:03d}", str(column), label, df[column])
            if text_insight:
                insights.append(text_insight)

    return pd.DataFrame(insights)
